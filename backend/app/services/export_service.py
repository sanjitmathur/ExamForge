"""Export generated papers to PDF and Word formats."""

import markdown
from io import BytesIO


def markdown_to_pdf(md_text: str, title: str = "Exam Paper") -> bytes:
    """Convert markdown to PDF using xhtml2pdf."""
    from xhtml2pdf import pisa

    html_content = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])

    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        body {{
            font-family: Helvetica, Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            color: #222;
        }}
        h1 {{ font-size: 17pt; text-align: center; margin-bottom: 8px; }}
        h2 {{ font-size: 13pt; margin-top: 18px; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
        h3 {{ font-size: 11pt; margin-top: 14px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        th, td {{ border: 1px solid #ccc; padding: 5px 8px; text-align: left; }}
        th {{ background: #f5f5f5; }}
        ol, ul {{ margin-left: 18px; }}
        p {{ margin: 5px 0; }}
    </style>
</head>
<body>{html_content}</body>
</html>"""

    buffer = BytesIO()
    pisa_status = pisa.CreatePDF(full_html, dest=buffer, encoding='utf-8')
    if pisa_status.err:
        raise RuntimeError("PDF generation failed")
    return buffer.getvalue()


def markdown_to_docx(md_text: str, title: str = "Exam Paper") -> bytes:
    """Convert markdown to DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = md_text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
            text = stripped.split(' ', 1)[1] if ' ' in stripped else stripped
            doc.add_paragraph(text, style='List Number')
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.bold = True
        elif stripped.startswith('---') or stripped.startswith('___'):
            doc.add_paragraph('_' * 50)
        else:
            doc.add_paragraph(stripped)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
